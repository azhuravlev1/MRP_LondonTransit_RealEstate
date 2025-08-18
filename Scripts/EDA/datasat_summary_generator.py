import os
from pathlib import Path
import random
import pandas as pd
from docx import Document
from docx.shared import Inches


ROOT_DIR = Path("/Users/andrey/Desktop/TMU/MRP/Code/Data")
OUTPUT_DOC = "dataset_summary.docx"
NUM_RODS_SAMPLE_YEARS = ['16', '17', '18', '19', '20', '21', '22', '23']  # Changed to match file naming format
NUMBAT_SAMPLE_YEARS = ['16', '17', '18', '19', '20', '21', '22', '23']


doc = Document()
doc.add_heading("Dataset Summary Report", 0)

# -------------------------
# PART 1: RODS Files
# -------------------------
rods_files = list(Path(ROOT_DIR).rglob("RODS OD/**/*.xls")) + \
             list(Path(ROOT_DIR).rglob("Rods OD 2000-2002/**/*.xls"))
doc.add_heading("RODS OD Summary", level=1)
doc.add_paragraph(
    f"There are {len(rods_files)} RODS files in the dataset, each containing 2 sheets: 'matrix' and 'zone'. "
    "The first 2 rows contain general file information, the next 2 define names of columns (the same names for all the files) and then the data goes. "
    f"For example, here is the header of a random RODS file ({Path(random.choice(rods_files)).name}):"
)

# Pick a random file and display first 4 rows from 'matrix'
sample_rods_file = random.choice(rods_files)
matrix_df = pd.read_excel(sample_rods_file, sheet_name="matrix", header=None, nrows=4)
doc.add_table(rows=0, cols=matrix_df.shape[1])
table = doc.tables[-1]
for _, row in matrix_df.iterrows():
    cells = table.add_row().cells
    for j, val in enumerate(row):
        cells[j].text = str(val)

doc.add_paragraph("Here is a sample of 10 rows from that file:")
sample_data = pd.read_excel(sample_rods_file, sheet_name="matrix", header=None, skiprows=4, nrows=10)
table = doc.add_table(rows=0, cols=sample_data.shape[1])
for _, row in sample_data.iterrows():
    cells = table.add_row().cells
    for j, val in enumerate(row):
        cells[j].text = str(val)

# Now document zone sheets

doc.add_paragraph("Now, here are contents of 'zone' sheets for each RODS file:")
for f in rods_files:
    doc.add_paragraph(f"File: {Path(f).name}", style='List Bullet')
    try:
        zone_df = pd.read_excel(f, sheet_name="zone")
        table = doc.add_table(rows=0, cols=zone_df.shape[1])
        for _, row in zone_df.iterrows():
            cells = table.add_row().cells
            for j, val in enumerate(row):
                cells[j].text = str(val)
    except Exception as e:
        doc.add_paragraph(f"Failed to parse zone sheet in {f}: {e}")

# -------------------------
# PART 2: NUMBAT Files
# -------------------------
doc.add_page_break()
doc.add_heading("NUMBAT Summary", level=1)
numbat_dir = ROOT_DIR / "NUMBAT"
print(f"Looking for NUMBAT files in: {numbat_dir}")
numbat_files = sorted([f for f in numbat_dir.glob("*.xlsx") if any(f"NBT{y}" in f.name for y in NUMBAT_SAMPLE_YEARS)])
print(f"Found {len(numbat_files)} NUMBAT files: {[f.name for f in numbat_files]}")

processed_years = set()

for f in numbat_files:
    year = next((y for y in NUMBAT_SAMPLE_YEARS if f"NBT{y}" in f.name), None)
    print(f"Processing file {f.name}, extracted year: {year}")
    if not year or year in processed_years:
        print(f"Skipping file {f.name} - year {year} already processed or invalid")
        continue

    processed_years.add(year)
    doc.add_heading(f"Sample NUMBAT File for {year}: {f.name}", level=2)

    try:
        print(f"Reading Excel file: {f}")
        xls = pd.ExcelFile(f)
        sheet_names = xls.sheet_names
        print(f"Found sheets: {sheet_names}")

        # Metadata sheet (assumed first)
        doc.add_paragraph(f"Metadata sheet: {sheet_names[0]}")
        meta_df = pd.read_excel(f, sheet_name=sheet_names[0])
        print(f"Metadata sheet shape: {meta_df.shape}")
        table = doc.add_table(rows=0, cols=meta_df.shape[1])
        for _, row in meta_df.iterrows():
            cells = table.add_row().cells
            for j, val in enumerate(row):
                cells[j].text = str(val)

        for sheet in sheet_names[1:]:
            doc.add_paragraph(f"\nSheet: {sheet}", style='List Bullet')
            try:
                header_rows = pd.read_excel(f, sheet_name=sheet, header=None, nrows=4)
                table = doc.add_table(rows=0, cols=header_rows.shape[1])
                for _, row in header_rows.iterrows():
                    cells = table.add_row().cells
                    for j, val in enumerate(row):
                        cells[j].text = str(val)

                doc.add_paragraph("Here is a sample of 10 rows from this sheet:")
                data_rows = pd.read_excel(f, sheet_name=sheet, header=None, skiprows=4, nrows=10)
                table = doc.add_table(rows=0, cols=data_rows.shape[1])
                for _, row in data_rows.iterrows():
                    cells = table.add_row().cells
                    for j, val in enumerate(row):
                        cells[j].text = str(val)
            except Exception as e:
                doc.add_paragraph(f"Failed to parse sheet {sheet}: {e}")

    except Exception as e:
        doc.add_paragraph(f"Failed to parse file {f.name}: {e}")


doc.save(OUTPUT_DOC)
print(f"✅ Summary saved to {OUTPUT_DOC}")
